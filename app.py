import os
import io
import json
import re
import time
import threading
from collections import defaultdict
from pathlib import Path
from typing import List, Dict, Tuple

import streamlit as st
import numpy as np
import faiss

from google.oauth2 import service_account
from vertexai import init as vertexai_init
from vertexai.language_models import TextEmbeddingModel
from vertexai.preview.generative_models import GenerativeModel, GenerationConfig, Part, Image

from docx import Document
from pypdf import PdfReader
import cv2
import pytesseract
from PIL import Image as PILImage

# --- constants ---------------------------------------------------------------

APP_DIR = Path(__file__).parent
DATA_DIR = APP_DIR / "data"
KB_DIR = APP_DIR / "kb"
EXTRACT_DIR = DATA_DIR / "kb_extracted"
INDEX_PATH = DATA_DIR / "faiss.index"
CORPUS_PATH = DATA_DIR / "corpus.json"

DATA_DIR.mkdir(parents=True, exist_ok=True)
KB_DIR.mkdir(parents=True, exist_ok=True)
EXTRACT_DIR.mkdir(parents=True, exist_ok=True)

CANDIDATE_MODELS = ["gemini-2.5-flash-lite", "gemini-2.5-pro"]
DEFAULT_LOCATION = "us-central1"

MAX_CONTEXT_TOKENS = 150_000
MAX_CHUNKS_INITIAL = 250
MAX_CHUNKS_FINAL = 15

MULTI_QUERY_VARIATIONS = 3
DEEP_RETRIEVAL_MULTIPLIER = 3
REQUIRED_UNIQUE_SOURCES = 6

MODEL_CONTEXT_LIMITS = {
    "gemini-2.5-flash-lite": 1_000_000,
    "gemini-2.5-pro": 2_000_000,
}

DEEP_RETRIEVAL_KEYWORDS = [
    "steps", "step-by-step", "process", "procedure", "workflow", "walk me through",
    "combine", "together", "multiple documents", "full instructions", "detailed",
    "all details", "comprehensive", "complete answer", "entire process",
    "different documents", "across", "multi-part", "split across",
]

# Performance optimization: limit images per document
MAX_IMAGES_PER_DOCUMENT = 5  
USE_VISION_MODEL_FOR_IMAGES = False  # Set to True to enable vision model, False for OCR only (faster)

# --- background polling ------------------------------------------------------

def check_db_changes():
    try:
        if check_kb_files_modified():
            print("KB files changed, triggering rebuild...")
            trigger_rebuild()
    except Exception as e:
        print(f"Error checking database changes: {e}")

def check_kb_files_modified() -> bool:
    try:
        last_check = st.session_state.get("last_kb_check", 0)
        current_time = time.time()
        if KB_DIR.exists():
            for file_path in KB_DIR.iterdir():
                if file_path.is_file() and file_path.stat().st_mtime > last_check:
                    st.session_state.last_kb_check = current_time
                    return True
        st.session_state.last_kb_check = current_time
        return False
    except Exception as e:
        print(f"Error checking file modifications: {e}")
        return False

def trigger_rebuild():
    try:
        st.session_state.kb_loaded = False
        st.session_state.kb_loading = True
        st.cache_resource.clear()
        st.rerun()
    except Exception as e:
        print(f"Error triggering rebuild: {e}")

def start_database_polling():
    def run_polling():
        while True:
            try:
                check_db_changes()
                time.sleep(300)
            except Exception as e:
                print(f"Polling error: {e}")
                time.sleep(300)
    threading.Thread(target=run_polling, daemon=True).start()

# --- token helpers -----------------------------------------------------------

def estimate_tokens(text: str) -> int:
    return len(text) // 4

def get_max_context_tokens(model_name: str) -> int:
    limit = MODEL_CONTEXT_LIMITS.get(model_name, MAX_CONTEXT_TOKENS)
    return int(limit * 0.8)

def truncate_to_token_limit(text: str, max_tokens: int) -> str:
    if estimate_tokens(text) <= max_tokens:
        return text
    return text[: max_tokens * 4] + "...[truncated]"

# --- text utilities ----------------------------------------------------------

def split_into_sentences(text: str) -> List[str]:
    sents = re.split(r"(?<=[\.\?\!])\s+", text.strip())
    return [s.strip() for s in sents if s.strip()]

def chunk_text(text: str, max_tokens: int = 200, overlap_sentences: int = 2) -> List[str]:
    sents = split_into_sentences(text)
    chunks, buf, token_est = [], [], 0
    for s in sents:
        s_tokens = max(1, len(s) // 4)
        if token_est + s_tokens > max_tokens and buf:
            chunks.append(" ".join(buf))
            buf = buf[-overlap_sentences:] if overlap_sentences > 0 else []
            token_est = sum(max(1, len(x)//4) for x in buf)
        buf.append(s)
        token_est += s_tokens
    if buf:
        chunks.append(" ".join(buf))
    validated = []
    for chunk in chunks:
        if estimate_tokens(chunk) > 2000:
            validated.extend(split_oversized_chunk(chunk, 2000))
        else:
            validated.append(chunk)
    return validated

def split_oversized_chunk(chunk: str, max_tokens: int = 2000) -> List[str]:
    words, sub_chunks = chunk.split(), []
    current, current_tokens = [], 0
    for word in words:
        word_tokens = len(word) // 4
        if current_tokens + word_tokens > max_tokens and current:
            sub_chunks.append(" ".join(current))
            current, current_tokens = [word], word_tokens
        else:
            current.append(word)
            current_tokens += word_tokens
    if current:
        sub_chunks.append(" ".join(current))
    return sub_chunks

# --- image extraction helpers -------------------------------------------------

def extract_images_from_pdf(pdf_bytes: bytes) -> List[bytes]:
    """Extract images from PDF pages."""
    images = []
    try:
        from pdf2image import convert_from_bytes
        pil_images = convert_from_bytes(pdf_bytes, dpi=200)
        for img in pil_images:
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            images.append(img_bytes.getvalue())
    except ImportError:
        # Fallback: try pypdf image extraction
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            for page_num, page in enumerate(reader.pages):
                if '/XObject' in page.get('/Resources', {}):
                    xobjects = page['/Resources']['/XObject'].get_object()
                    for obj_name in xobjects:
                        obj = xobjects[obj_name]
                        if obj.get('/Subtype') == '/Image':
                            try:
                                img_data = obj.get_data()
                                images.append(img_data)
                            except:
                                pass
        except:
            pass
    except Exception:
        pass
    return images

def extract_images_from_docx(docx_bytes: bytes) -> List[bytes]:
    """Extract images from DOCX file."""
    images = []
    try:
        import zipfile
        docx_zip = zipfile.ZipFile(io.BytesIO(docx_bytes))
        # Images in DOCX are stored in word/media/
        for file_info in docx_zip.filelist:
            if file_info.filename.startswith('word/media/'):
                image_data = docx_zip.read(file_info.filename)
                # Filter common image formats
                if len(image_data) > 4 and any(image_data[:4] == sig for sig in [b'\x89PNG', b'\xff\xd8\xff', b'GIF8']):
                    images.append(image_data)
                elif len(image_data) > 12 and b'WEBP' in image_data[:12]:
                    images.append(image_data)
        docx_zip.close()
    except Exception:
        pass
    return images

def extract_images_from_doc(doc_bytes: bytes) -> List[bytes]:
    """Extract images from .doc (OLE) files."""
    images = []
    try:
        import olefile
        file_obj = io.BytesIO(doc_bytes)
        if olefile.isOleFile(file_obj):
            ole = olefile.OleFileIO(file_obj)
            # Images in .doc files can be in various streams
            for stream_name in ole.listdir():
                if stream_name and len(stream_name) > 0:
                    stream_path = stream_name[0] if isinstance(stream_name, list) else stream_name
                    try:
                        stream = ole.openstream(stream_path)
                        data = stream.read()
                        # Check if it looks like an image (PNG, JPEG signatures)
                        if len(data) > 10:
                            if data[:4] == b'\x89PNG' or data[:2] == b'\xff\xd8':
                                images.append(data)
                            elif data[:3] == b'GIF':
                                images.append(data)
                    except:
                        pass
            ole.close()
    except ImportError:
        pass
    except Exception:
        pass
    return images

def process_image_with_ocr_or_vision(image_bytes: bytes, model_name: str, project_id: str, location: str, credentials, silent: bool = False) -> str:
    """Process image through OCR first, then vision model if needed."""
    # Try OCR first (faster, cheaper)
    try:
        img = PILImage.open(io.BytesIO(image_bytes)).convert("RGB")
        arr = np.array(img)[:, :, ::-1]
        gray = cv2.cvtColor(arr, cv2.COLOR_BGR2GRAY)
        gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
        ocr_text = pytesseract.image_to_string(gray) or ""
        
        # If OCR got substantial text, use it
        if len(ocr_text.strip()) > 20:
            return f"[Image Content via OCR]\n{ocr_text.strip()}"
    except Exception:
        pass
    
    # Fallback to vision model only if enabled (slow/expensive)
    if USE_VISION_MODEL_FOR_IMAGES and project_id and credentials:
        try:
            vertexai_init(project=project_id, location=location, credentials=credentials)
            model = GenerativeModel(model_name)
            
            mime_type = "image/png"
            if image_bytes.startswith(b"\xff\xd8\xff"):
                mime_type = "image/jpeg"
            elif image_bytes.startswith(b"GIF8"):
                mime_type = "image/gif"
            
            image_part = Part.from_data(image_bytes, mime_type=mime_type)
            prompt = "Extract all text and describe any important visual elements, screenshots, diagrams, or procedures shown in this HBS NetView image. Be thorough and include all details visible."
            
            response = model.generate_content([prompt, image_part])
            if response.text and len(response.text.strip()) > 20:
                return f"[Image Content via Vision Model]\n{response.text.strip()}"
        except Exception:
            pass
    
    return ""

# --- extraction helpers ------------------------------------------------------

def extract_text_from_docx_bytes(b: bytes) -> str:
    """Extract text AND images from DOCX."""
    text_parts = []
    
    # Extract text
    try:
        doc = Document(io.BytesIO(b))
        docx_text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        if docx_text.strip():
            text_parts.append(docx_text)
    except Exception:
        pass
    
    # Extract and process images (limited for performance)
    try:
        images = extract_images_from_docx(b)
        if images:
            # Limit to MAX_IMAGES_PER_DOCUMENT for faster processing
            for i, img_bytes in enumerate(images[:MAX_IMAGES_PER_DOCUMENT]):
                # Get credentials from session state if available
                model_name = getattr(st.session_state, 'model_name', None) or CANDIDATE_MODELS[0]
                project_id = getattr(st.session_state, 'project_id', None)
                location = getattr(st.session_state, 'location', None) or DEFAULT_LOCATION
                creds = getattr(st.session_state, 'creds', None)
                
                if project_id and creds:
                    img_text = process_image_with_ocr_or_vision(
                        img_bytes,
                        model_name,
                        project_id,
                        location,
                        creds,
                        silent=True
                    )
                    if img_text:
                        text_parts.append(img_text)
    except Exception as e:
        print(f"Error processing DOCX images: {e}")
    
    return "\n\n".join(text_parts) if text_parts else ""

def extract_text_from_doc_bytes(b: bytes) -> str:
    """Extract text AND images from .doc files using multiple fallback methods."""
    text_parts = []
    errors = []
    
    # Method 1: Try python-docx (sometimes works for older .doc if converted)
    try:
        from docx import Document
        doc = Document(io.BytesIO(b))
        text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        if text.strip():
            text_parts.append(text)
            print(f"✓ Successfully extracted text from .doc using python-docx: {len(text)} characters")
    except Exception as e:
        errors.append(f"python-docx: {str(e)[:100]}")
    
    # Method 2: Try textract if available
    if not text_parts:
        try:
            import textract
            text = textract.process(io.BytesIO(b), extension="doc").decode("utf-8").strip()
            if text.strip():
                text_parts.append(text)
                print(f"✓ Successfully extracted text from .doc using textract: {len(text)} characters")
        except ImportError:
            errors.append("textract: not installed")
        except Exception as e:
            errors.append(f"textract: {str(e)[:100]}")
    
    # Method 3: Try antiword command line tool
    if not text_parts:
        try:
            import subprocess, tempfile
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_file:
                tmp_file.write(b)
                tmp_path = tmp_file.name
            try:
                result = subprocess.run(["antiword", tmp_path], capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    text_parts.append(result.stdout.strip())
                    print(f"✓ Successfully extracted text from .doc using antiword: {len(result.stdout)} characters")
                elif result.returncode != 0:
                    errors.append(f"antiword: exit code {result.returncode}")
            finally:
                try:
                    os.unlink(tmp_path)
                except FileNotFoundError:
                    pass
        except FileNotFoundError:
            errors.append("antiword: command not found (install: apt-get install antiword or brew install antiword)")
        except Exception as e:
            errors.append(f"antiword: {str(e)[:100]}")
    
    # Method 4: Try catdoc if available
    if not text_parts:
        try:
            import subprocess, tempfile
            with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_file:
                tmp_file.write(b)
                tmp_path = tmp_file.name
            try:
                result = subprocess.run(["catdoc", tmp_path], capture_output=True, text=True, timeout=30)
                if result.returncode == 0 and result.stdout.strip():
                    text_parts.append(result.stdout.strip())
                    print(f"✓ Successfully extracted text from .doc using catdoc: {len(result.stdout)} characters")
                elif result.returncode != 0:
                    errors.append(f"catdoc: exit code {result.returncode}")
            finally:
                try:
                    os.unlink(tmp_path)
                except FileNotFoundError:
                    pass
        except FileNotFoundError:
            errors.append("catdoc: command not found (install: apt-get install catdoc or brew install catdoc)")
        except Exception as e:
            errors.append(f"catdoc: {str(e)[:100]}")
    
    # Method 5: Try olefile for structured storage (better for .doc format)
    if not text_parts:
        try:
            import olefile
            file_obj = io.BytesIO(b)
            if olefile.isOleFile(file_obj):
                ole = olefile.OleFileIO(file_obj)
                streams = ole.listdir()
                if ole.exists('WordDocument'):
                    stream = ole.openstream('WordDocument')
                    data = stream.read()
                    # Better text extraction from Word binary format
                    text = ""
                    # Look for readable text sequences
                    for i in range(len(data) - 1):
                        byte1 = data[i]
                        # ASCII printable range
                        if 32 <= byte1 < 127:
                            text += chr(byte1)
                        elif byte1 in (10, 13, 9):  # newline, carriage return, tab
                            text += ' ' if byte1 != 10 else '\n'
                    
                    # Clean up excessive whitespace
                    text = re.sub(r'\s+', ' ', text).strip()
                    ole.close()
                    if text.strip() and len(text.strip()) > 50:
                        text_parts.append(text.strip())
                        print(f"✓ Successfully extracted text from .doc using olefile: {len(text.strip())} characters")
                else:
                    available_streams = [str(s) for s in streams[:5]]
                    errors.append(f"olefile: WordDocument stream not found. Available: {available_streams}")
                    ole.close()
            else:
                errors.append("olefile: Not a valid OLE file")
        except ImportError:
            errors.append("olefile: not installed (pip install olefile)")
        except Exception as e:
            errors.append(f"olefile: {str(e)[:100]}")
    
    # Method 6: Last resort - try to extract any readable text
    if not text_parts:
        try:
            text = b.decode("utf-8", errors="ignore")
            readable = "".join(ch for ch in text if ch.isprintable() or ch in "\n\r\t").strip()
            # Only use if we got substantial readable text (not just random characters)
            if len(readable) > 200 and len([c for c in readable if c.isalpha()]) > 50:
                text_parts.append(readable)
                print(f"✓ Successfully extracted text from .doc using raw extraction: {len(readable)} characters")
            else:
                errors.append("raw extraction: insufficient readable text")
        except Exception as e:
            errors.append(f"raw extraction: {str(e)[:100]}")
    
    # Extract and process images from .doc file (limited for performance)
    try:
        images = extract_images_from_doc(b)
        if images:
            # Limit to MAX_IMAGES_PER_DOCUMENT
            for img_bytes in images[:MAX_IMAGES_PER_DOCUMENT]:
                # Get credentials from session state if available
                model_name = getattr(st.session_state, 'model_name', None) or CANDIDATE_MODELS[0]
                project_id = getattr(st.session_state, 'project_id', None)
                location = getattr(st.session_state, 'location', None) or DEFAULT_LOCATION
                creds = getattr(st.session_state, 'creds', None)
                
                if project_id and creds:
                    img_text = process_image_with_ocr_or_vision(
                        img_bytes,
                        model_name,
                        project_id,
                        location,
                        creds,
                        silent=True
                    )
                    if img_text:
                        text_parts.append(img_text)
    except Exception as e:
        errors.append(f"image extraction: {str(e)[:100]}")
    
    # Return combined text and image content
    if text_parts:
        return "\n\n".join(text_parts)
    
    # If all methods fail, log errors (only show first few to avoid spam)
    if errors and len(errors) > 0:
        print(f"✗ FAILED to extract text from .doc file")
        print(f"  Errors: {'; '.join(errors[:3])}")  # Show first 3 errors only
    
    return ""

def extract_text_from_pdf_bytes(b: bytes) -> str:
    """Extract text AND images from PDF."""
    text_parts = []
    
    # Extract text
    try:
        reader = PdfReader(io.BytesIO(b))
        pdf_text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
        if pdf_text.strip():
            text_parts.append(pdf_text)
    except Exception:
        pass
    
    # Extract and process images (limited for performance)
    try:
        images = extract_images_from_pdf(b)
        if images:
            # Limit to MAX_IMAGES_PER_DOCUMENT for faster processing
            for img_bytes in images[:MAX_IMAGES_PER_DOCUMENT]:
                # Get credentials from session state if available
                model_name = getattr(st.session_state, 'model_name', None) or CANDIDATE_MODELS[0]
                project_id = getattr(st.session_state, 'project_id', None)
                location = getattr(st.session_state, 'location', None) or DEFAULT_LOCATION
                creds = getattr(st.session_state, 'creds', None)
                
                if project_id and creds:
                    img_text = process_image_with_ocr_or_vision(
                        img_bytes,
                        model_name,
                        project_id,
                        location,
                        creds,
                        silent=True
                    )
                    if img_text:
                        text_parts.append(img_text)
    except Exception as e:
        print(f"Error processing PDF images: {e}")
    
    return "\n\n".join(text_parts) if text_parts else ""

def extract_text_from_image_bytes(b: bytes) -> str:
    try:
        img = PILImage.open(io.BytesIO(b)).convert("RGB")
        arr = np.array(img)[:, :, ::-1]
        gray = cv2.cvtColor(arr, cv2.COLOR_BGR2GRAY)
        gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
        return (pytesseract.image_to_string(gray) or "").strip()
    except Exception:
        return ""

# --- structured OCR parsing --------------------------------------------------

def parse_report_data_from_ocr(ocr_text: str, filename: str) -> List[Dict]:
    lines = [line.strip() for line in ocr_text.split("\n") if line.strip()]
    if "overdue" in filename.lower() or "overdue" in ocr_text.lower():
        return parse_overdue_report(lines, filename)
    if "outbound" in filename.lower() or "outbound" in ocr_text.lower():
        return parse_outbound_report(lines, filename)
    if "equipment list" in filename.lower() or "equipment list" in ocr_text.lower():
        return parse_equipment_list_report(lines, filename)
    return []

def parse_overdue_report(lines: List[str], filename: str) -> List[Dict]:
    data = []
    for i, line in enumerate(lines):
        if re.match(r"^[A-Z\s]+$", line) and len(line) > 3:
            if i + 1 >= len(lines):
                continue
            next_line = lines[i + 1]
            contract_match = re.search(r"C\d+R", next_line)
            if not contract_match:
                continue
            customer_name = line
            contract = contract_match.group()
            phone = stock = make = model = equipment_type = ""
            year = serial = date_out = expected = days_over = ""
            for j in range(i, min(i + 5, len(lines))):
                current_line = lines[j]
                phone_match = re.search(r"\(\d{3}\)\s*\d{3}-\d{4}", current_line)
                if phone_match:
                    phone = phone_match.group()
                stock_match = re.search(r"\b\d{5}\b", current_line)
                if stock_match:
                    stock = stock_match.group()
                make_match = re.search(r"\b(BOB|KUB|JD|BOM)\b", current_line)
                if make_match:
                    make = make_match.group()
                model_match = re.search(
                    r"\b(T650|E32|E42|U55-4R3AP|U35-4R3A|E26|KX080R3AT3|KX121R3TA|KX121RRATS|211D-50|690B|442)\b",
                    current_line,
                )
                if model_match:
                    model = model_match.group()
                type_match = re.search(r"\b(SKIDSTEER|EXCAVATOR|ROLLER)\b", current_line)
                if type_match:
                    equipment_type = type_match.group()
                year_match = re.search(r"\b(2013|2014|2015|2016|1979|2006|2008|2012)\b", current_line)
                if year_match:
                    year = year_match.group()
                serial_match = re.search(r"\b[A-Z0-9]{6,}\b", current_line)
                if serial_match and len(serial_match.group()) > 6:
                    serial = serial_match.group()
                date_match = re.search(r"\b\d{2}/\d{2}/\d{4}\b", current_line)
                if date_match:
                    if not date_out:
                        date_out = date_match.group()
                    else:
                        expected = date_match.group()
                days_match = re.search(r"\b\d{1,4}\b", current_line)
                if days_match and days_match.group().isdigit():
                    days_over = days_match.group()
            if customer_name and contract:
                data.append(
                    {
                        "customer_name": customer_name,
                        "contract": contract,
                        "phone": phone,
                        "stock_number": stock,
                        "make": make,
                        "model": model,
                        "equipment_type": equipment_type,
                        "year": year,
                        "serial": serial,
                        "date_out": date_out,
                        "expected_due": expected,
                        "days_overdue": days_over,
                        "source": filename,
                        "report_type": "Overdue Equipment Report",
                    }
                )
    return data

def parse_outbound_report(lines: List[str], filename: str) -> List[Dict]:
    data = []
    for i, line in enumerate(lines):
        if re.match(r"^[A-Z\s]+$", line) and len(line) > 3:
            if i + 1 >= len(lines):
                continue
            next_line = lines[i + 1]
            contract_match = re.search(r"C\d+R", next_line)
            if not contract_match:
                continue
            customer_name = line
            contract = contract_match.group()
            phone = stock = make = model = equipment_type = ""
            year = serial = date_time_out = ""
            for j in range(i, min(i + 5, len(lines))):
                current_line = lines[j]
                phone_match = re.search(r"\(\d{3}\)\s*\d{3}-\d{4}", current_line)
                if phone_match:
                    phone = phone_match.group()
                stock_match = re.search(r"\b\d{5}\b", current_line)
                if stock_match:
                    stock = stock_match.group()
                make_match = re.search(r"\b(BOB|KUB|JD|BOM)\b", current_line)
                if make_match:
                    make = make_match.group()
                model_match = re.search(
                    r"\b(T650|E32|E42|U55-4R3AP|U35-4R3A|E26|KX080R3AT3|KX121R3TA|KX121RRATS|211D-50|690B|442)\b",
                    current_line,
                )
                if model_match:
                    model = model_match.group()
                type_match = re.search(r"\b(SKIDSTEER|EXCAVATOR|ROLLER)\b", current_line)
                if type_match:
                    equipment_type = type_match.group()
                year_match = re.search(r"\b(2013|2014|2015|2016|1979|2006|2008|2012)\b", current_line)
                if year_match:
                    year = year_match.group()
                serial_match = re.search(r"\b[A-Z0-9]{6,}\b", current_line)
                if serial_match and len(serial_match.group()) > 6:
                    serial = serial_match.group()
                datetime_match = re.search(r"\b\d{2}/\d{2}/\d{4}\s+\d{2}:\d{2}\s+[AP]M\b", current_line)
                if datetime_match:
                    date_time_out = datetime_match.group()
            if customer_name and contract:
                data.append(
                    {
                        "customer_name": customer_name,
                        "contract": contract,
                        "phone": phone,
                        "stock_number": stock,
                        "make": make,
                        "model": model,
                        "equipment_type": equipment_type,
                        "year": year,
                        "serial": serial,
                        "date_time_out": date_time_out,
                        "source": filename,
                        "report_type": "Rental Outbound Report",
                    }
                )
    return data

def parse_equipment_list_report(lines: List[str], filename: str) -> List[Dict]:
    data = []
    for line in lines:
        stock_match = re.search(r"\b\d{5}\b", line)
        if not stock_match:
            continue
        stock = stock_match.group()
        make_match = re.search(r"\b(BOB|KUB|JD|BOM)\b", line)
        model_match = re.search(
            r"\b(T650|E32|E42|U55-4R3AP|U35-4R3A|E26|KX080R3AT3|KX121R3TA|KX121RRATS|211D-50|690B|442)\b",
            line,
        )
        type_match = re.search(r"\b(SKIDSTEER|EXCAVATOR|ROLLER)\b", line)
        year_match = re.search(r"\b(2013|2014|2015|2016|1979|2006|2008|2012)\b", line)
        serial_match = re.search(r"\b[A-Z0-9]{6,}\b", line)
        meter_match = re.search(r"\b\d+\b", line)
        data.append(
            {
                "stock_number": stock,
                "make": make_match.group() if make_match else "",
                "model": model_match.group() if model_match else "",
                "equipment_type": type_match.group() if type_match else "",
                "year": year_match.group() if year_match else "",
                "serial": serial_match.group() if serial_match and len(serial_match.group()) > 6 else "",
                "location": "",
                "meter": meter_match.group() if meter_match else "",
                "source": filename,
                "report_type": "Rental Equipment List",
            }
        )
    return data

# --- embeddings/index --------------------------------------------------------

def embed_texts(texts: List[str], project_id: str, location: str, credentials, silent: bool = False) -> np.ndarray:
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = TextEmbeddingModel.from_pretrained("text-embedding-005")

        all_embeddings: List[np.ndarray] = []
        valid_texts, skipped = [], 0

        for i, text in enumerate(texts):
            if estimate_tokens(text) > 10_000:
                skipped += 1
                if not silent and skipped <= 5:
                    st.warning(f"Skipping text {i+1} (>{10_000} tokens)")
                all_embeddings.append(np.zeros(768))
            else:
                valid_texts.append((i, text))

        if not silent and skipped > 5:
            st.warning(f"Skipped {skipped} texts due to size limits")

        MAX_BATCH_TOKENS = 15_000
        current_batch, current_tokens = [], 0

        def flush_batch():
            nonlocal current_batch, current_tokens
            batch_texts = [item[1] for item in current_batch]
            try:
                results = model.get_embeddings(batch_texts)
                for j, embedding in enumerate(results):
                    orig_idx = current_batch[j][0]
                    while len(all_embeddings) < orig_idx:
                        all_embeddings.append(np.zeros(768))
                    all_embeddings.append(embedding.values)
            except Exception as batch_error:
                if not silent:
                    st.error(f"Embedding batch error: {batch_error}")
                for orig_idx, _ in current_batch:
                    while len(all_embeddings) < orig_idx:
                        all_embeddings.append(np.zeros(768))
                    all_embeddings.append(np.zeros(768))
            current_batch, current_tokens = [], 0

        for original_idx, text in valid_texts:
            tokens = estimate_tokens(text)
            if current_batch and (current_tokens + tokens > MAX_BATCH_TOKENS or len(current_batch) >= 100):
                flush_batch()
            current_batch.append((original_idx, text))
            current_tokens += tokens

        if current_batch:
            flush_batch()

        return np.array(all_embeddings).astype(np.float32)
    except Exception as e:
        if not silent:
            st.error(f"Embedding error: {e}")
        return np.array([])

def build_faiss_index(corpus: List[Dict], project_id: str, location: str, credentials, silent: bool = False) -> Tuple[faiss.IndexFlatIP, List[Dict]]:
    if not corpus:
        return None, []
    texts = [item["text"] for item in corpus]
    embeddings = embed_texts(texts, project_id, location, credentials, silent=silent)
    if embeddings.size == 0:
        return None, []
    index = faiss.IndexFlatIP(embeddings.shape[1])
    faiss.normalize_L2(embeddings)
    index.add(embeddings)
    return index, corpus

# --- retrieval helpers -------------------------------------------------------

def expand_query(query: str) -> str:
    q = query.lower()
    if "overdue" in q:
        return f"{query} overdue equipment report rental"
    if "outbound" in q:
        return f"{query} outbound report rental equipment"
    if "equipment" in q:
        return f"{query} equipment list rental"
    if "customer" in q:
        return f"{query} customer contract phone"
    if "stock" in q:
        return f"{query} stock number equipment"
    if "serial" in q:
        return f"{query} serial number equipment"
    return query

def decompose_complex_query(query: str, model_name: str, project_id: str, location: str, credentials) -> List[str]:
    """Break down complex queries into sub-queries to find different aspects across documents."""
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = GenerativeModel(model_name)
        
        prompt = f"""Given this question about HBS NetView, break it into {MULTI_QUERY_VARIATIONS + 2} separate search queries that target different aspects, documents, or steps.
Each sub-query should use different terminology that might appear in different documents.

Original question: {query}

Return ONLY a JSON array of strings, each targeting a different aspect, document type, procedure step, or data source.
Be specific about different procedures, documents, screens, processes, or related concepts.
"""
        response = model.generate_content(
            prompt,
            generation_config=GenerationConfig(temperature=0.6, max_output_tokens=400),
        )
        
        if response.text:
            sub_queries = json.loads(response.text.strip())
            if isinstance(sub_queries, list) and len(sub_queries) > 1:
                return [query] + [q for q in sub_queries if isinstance(q, str) and q.strip()][:MULTI_QUERY_VARIATIONS + 2]
    except Exception as e:
        print(f"Query decomposition error: {e}")
    return [query]

def generate_query_variations(query: str, model_name: str, project_id: str, location: str, credentials) -> List[str]:
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = GenerativeModel(model_name)
        prompt = f"""Generate {MULTI_QUERY_VARIATIONS} alternative search queries that might fetch different HBS NetView documents.

Original query: {query}

Return ONLY a JSON array of strings.
"""
        response = model.generate_content(
            prompt,
            generation_config=GenerationConfig(temperature=0.5, max_output_tokens=200),
        )
        if response.text:
            variations = json.loads(response.text.strip())
            if isinstance(variations, list):
                variations = [v for v in variations if isinstance(v, str) and v.strip()]
                return [query] + variations[:MULTI_QUERY_VARIATIONS]
    except Exception:
        pass
    return [query]

def needs_deep_retrieval(query: str) -> bool:
    """Detect queries that need information from multiple documents."""
    lowered = query.lower()
    
    # Existing keywords
    if any(keyword in lowered for keyword in DEEP_RETRIEVAL_KEYWORDS):
        return True
    
    # Additional patterns for multi-document queries
    multi_doc_patterns = [
        r"across\s+(multiple|several|different|various)",
        r"(both|all|each|every)\s+\w+\s+and",
        r"compare",
        r"difference",
        r"relation",
        r"related to",
        r"following.*and",
        r"multiple steps",
        r"different.*documents?",
        r"several.*sources?",
        r"step.*step",
        r"\d+\s+and\s+\d+",
    ]
    
    if any(re.search(pattern, lowered) for pattern in multi_doc_patterns):
        return True
    
    # If question has multiple parts
    clause_markers = r"\s+(and|then|after|before|when|also|plus)\s+"
    if len(re.split(clause_markers, lowered)) > 2:
        return True
    
    return False

def find_connected_documents(seed_chunks: List[Dict], corpus: List[Dict], 
                             project_id: str, location: str, credentials, index) -> List[Dict]:
    """Find documents that share entities/concepts with already-retrieved chunks."""
    if not seed_chunks or len(seed_chunks) == 0 or index is None:
        return []
    
    try:
        # Extract key terms from seed chunks
        seed_texts = [chunk["text"][:500] for chunk in seed_chunks[:10]]
        
        # Get embeddings for seed chunks
        query_vec = embed_texts(seed_texts, project_id, location, credentials, silent=True)
        if query_vec.size == 0:
            return []
        
        # Average embeddings to find related documents
        avg_query = np.mean(query_vec, axis=0).reshape(1, -1)
        faiss.normalize_L2(avg_query)
        
        k_connected = min(50, len(corpus))
        scores, indices = index.search(avg_query, k_connected)
        
        connected = []
        seen_sources = {chunk["source"] for chunk in seed_chunks}
        
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(corpus) and score >= 0.15:
                candidate = corpus[idx]
                if candidate["source"] not in seen_sources or len(seen_sources) < 5:
                    connected.append({
                        **candidate,
                        "similarity_score": float(score),
                        "query_source": "connected_doc",
                        "connection_type": "semantic_bridge"
                    })
                    seen_sources.add(candidate["source"])
        
        return connected
    except Exception as e:
        print(f"Error finding connected documents: {e}")
    return []

def diversify_chunks(candidates: List[Dict], lambda_param: float, top_k: int) -> List[Dict]:
    """MMR-style diversification with better source balancing."""
    if not candidates:
        return []
    
    selected: List[Dict] = []
    used_sources = defaultdict(int)
    pool = candidates[:]
    
    # Calculate max chunks per source based on diversity of sources
    unique_sources = len(set(c["source"] for c in candidates[:50]))
    max_per_source = max(1, top_k // max(3, unique_sources // 2))
    
    while pool and len(selected) < top_k:
        best_idx, best_score, best_chunk = None, -float("inf"), None
        
        for idx, chunk in enumerate(pool):
            base_score = chunk.get("similarity_score", 0.0)
            
            # Penalty for over-represented sources
            source_count = used_sources[chunk["source"]]
            if source_count >= max_per_source:
                penalty = (1 - lambda_param) * 1.0
            elif chunk["source"] in used_sources:
                penalty = (1 - lambda_param) * 0.4 * (source_count / max_per_source)
            else:
                penalty = 0.0
            
            score = lambda_param * base_score - penalty
            
            if score > best_score:
                best_idx, best_score, best_chunk = idx, score, chunk
        
        if best_chunk is None:
            break
        
        selected.append(best_chunk)
        used_sources[best_chunk["source"]] += 1
        pool.pop(best_idx)
    
    return selected

def ensure_source_diversity(chunks: List[Dict], candidates: List[Dict], min_sources: int) -> List[Dict]:
    selected_sources = {chunk["source"] for chunk in chunks}
    if len(selected_sources) >= min_sources:
        return chunks
    for chunk in candidates:
        if chunk["source"] not in selected_sources:
            chunks.append(chunk)
            selected_sources.add(chunk["source"])
        if len(selected_sources) >= min_sources:
            break
    return chunks

def cluster_chunks(chunks: List[Dict]) -> List[Dict]:
    if not chunks:
        return []
    grouped = defaultdict(list)
    for chunk in chunks:
        grouped[chunk.get("source", "Unknown")].append(chunk)
    merged = []
    for source, source_chunks in grouped.items():
        source_chunks.sort(key=lambda x: str(x.get("chunk_id")))
        merged_text = "\n".join(ch["text"] for ch in source_chunks)
        merged.append(
            {
                "source": source,
                "text": merged_text,
                "similarity_score": max(ch.get("similarity_score", 0) for ch in source_chunks),
                "rerank_score": max(ch.get("rerank_score", 0) for ch in source_chunks),
                "chunk_id": source_chunks[0].get("chunk_id"),
            }
        )
    merged.sort(key=lambda x: (x.get("rerank_score", 0), x.get("similarity_score", 0)), reverse=True)
    return merged

def verify_answer_quality(query: str, answer: str, context_chunks: List[Dict], model_name: str, project_id: str, location: str, credentials) -> Dict:
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = GenerativeModel(model_name)
        context_summary = "\n".join(f"- {chunk['text'][:300]}" for chunk in context_chunks[:5])
        prompt = f"""Verify answer quality with context support.

QUESTION:
{query}

ANSWER:
{answer}

CONTEXT:
{context_summary}

Return JSON: {{"relevance":8,"support":7,"completeness":6,"accuracy_risk":2}}
"""
        response = model.generate_content(
            prompt,
            generation_config=GenerationConfig(temperature=0.1, max_output_tokens=200),
        )
        if response.text:
            result = json.loads(response.text.strip())
            if isinstance(result, dict):
                result["needs_improvement"] = result.get("support", 0) < 5 or result.get("accuracy_risk", 10) > 6
                return result
    except Exception:
        pass
    return {"relevance": 6, "support": 6, "completeness": 6, "accuracy_risk": 4, "needs_improvement": False}

def rerank_chunks(query: str, chunks: List[Dict], model_name: str, project_id: str, location: str, credentials, top_k: int) -> List[Dict]:
    if len(chunks) <= top_k:
        return chunks
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = GenerativeModel(model_name)
        snippet = "\n".join(f"[{i}] Source: {chunk.get('source', 'Unknown')}\n{chunk['text'][:600]}" for i, chunk in enumerate(chunks[:50]))
        prompt = f"""Score chunk relevance to query on 0-10 scale. Higher scores for chunks that help answer the complete question.

QUERY: {query}

CHUNKS:
{snippet}

Return JSON array of scores [score1, score2, ...] matching the order of chunks.
Consider: direct relevance, completeness of information, and whether it complements other chunks.
"""
        response = model.generate_content(
            prompt,
            generation_config=GenerationConfig(temperature=0.1, max_output_tokens=500),
        )
        if response.text:
            scores = json.loads(response.text.strip())
            if isinstance(scores, list):
                for i, score in enumerate(scores[: len(chunks)]):
                    chunks[i]["rerank_score"] = float(score)
                chunks.sort(key=lambda x: (x.get("rerank_score", 0), x.get("similarity_score", 0)), reverse=True)
                return chunks[:top_k]
    except Exception:
        pass
    return chunks[:top_k]

def build_optimized_context(chunks: List[Dict], max_tokens: int) -> str:
    if not chunks:
        return "No relevant information found in knowledge base."
    seen = set()
    unique = []
    for chunk in chunks:
        signature = (chunk.get("source", ""), chunk["text"][:200])
        if signature not in seen:
            seen.add(signature)
            unique.append(chunk)
    unique.sort(key=lambda x: (x.get("rerank_score", 0), x.get("similarity_score", 0)), reverse=True)
    context_parts, current_tokens, current_source = [], 0, None
    for chunk in unique:
        source = chunk.get("source", "Unknown")
        chunk_text = f"Source: {source}\nContent: {chunk['text']}\n"
        chunk_tokens = estimate_tokens(chunk_text)
        if current_tokens + chunk_tokens > max_tokens:
            remaining = max_tokens - current_tokens
            if remaining > 100:
                truncated = truncate_to_token_limit(chunk["text"], remaining - 50)
                context_parts.append(f"Source: {source}\nContent: {truncated}\n")
            break
        if source != current_source:
            if current_source is not None:
                context_parts.append("")
            current_source = source
        context_parts.append(chunk_text)
        current_tokens += chunk_tokens
    return "\n".join(context_parts)

def search_index(query: str, index, corpus: List[Dict], project_id: str, location: str, credentials, model_name: str, deep_mode: bool, k: int, min_similarity: float = 0.2) -> List[Dict]:
    if index is None or not corpus:
        return []
    
    try:
        # For deep mode, use query decomposition instead of just variations
        if deep_mode:
            queries = decompose_complex_query(query, model_name, project_id, location, credentials)
        else:
            queries = generate_query_variations(query, model_name, project_id, location, credentials)
        
        all_candidates, seen_indices = [], set()
        
        # PASS 1: Initial retrieval with more candidates for deep mode
        base_k = min(MAX_CHUNKS_INITIAL, len(corpus))
        initial_k = min(base_k * (DEEP_RETRIEVAL_MULTIPLIER * 2 if deep_mode else 1), len(corpus))
        
        # Lower threshold for deep mode to catch more potentially relevant docs
        similarity_threshold = min_similarity * 0.7 if deep_mode else min_similarity
        
        for query_variant in queries:
            expanded = expand_query(query_variant)
            query_vec = embed_texts([expanded], project_id, location, credentials, silent=True)
            
            if query_vec.size == 0:
                continue
            
            faiss.normalize_L2(query_vec)
            scores, indices = index.search(query_vec, initial_k)
            
            for score, idx in zip(scores[0], indices[0]):
                if idx < len(corpus) and score >= similarity_threshold and idx not in seen_indices:
                    seen_indices.add(idx)
                    all_candidates.append({
                        **corpus[idx], 
                        "similarity_score": float(score), 
                        "query_source": query_variant
                    })
        
        if not all_candidates:
            return []
        
        # PASS 2: For deep mode, find documents connected to already-retrieved ones
        if deep_mode and len(all_candidates) > 0:
            connected_candidates = find_connected_documents(
                all_candidates[:20],
                corpus,
                project_id,
                location,
                credentials,
                index
            )
            
            # Add connected candidates (avoid duplicates)
            existing_sigs = {(c["source"], str(c.get("chunk_id", ""))) for c in all_candidates}
            for conn in connected_candidates:
                sig = (conn["source"], str(conn.get("chunk_id", "")))
                if sig not in existing_sigs:
                    all_candidates.append(conn)
                    existing_sigs.add(sig)
        
        # Sort by similarity
        all_candidates.sort(key=lambda x: x["similarity_score"], reverse=True)
        
        # Diversify with more emphasis on source diversity in deep mode
        lambda_param = 0.6 if deep_mode else 0.75
        diversified_k = k * (DEEP_RETRIEVAL_MULTIPLIER * 2 if deep_mode else 1)
        
        diversified = diversify_chunks(
            all_candidates,
            lambda_param=lambda_param,
            top_k=diversified_k,
        )
        
        # Rerank but keep more candidates
        rerank_k = len(diversified) if deep_mode else min(k * 2, len(diversified))
        reranked = rerank_chunks(
            query,
            diversified,
            model_name,
            project_id,
            location,
            credentials,
            top_k=rerank_k,
        )
        
        # Ensure source diversity with higher minimum for deep mode
        min_sources = REQUIRED_UNIQUE_SOURCES * 2 if deep_mode else 1
        final = ensure_source_diversity(
            reranked, 
            diversified, 
            min_sources=min_sources
        )
        
        # Return more chunks for deep mode
        return final[:k * (DEEP_RETRIEVAL_MULTIPLIER if deep_mode else 1)]
        
    except Exception as e:
        st.error(f"Search error: {e}")
        import traceback
        traceback.print_exc()
        return []

# --- persistence -------------------------------------------------------------

def load_index_and_corpus():
    try:
        if INDEX_PATH.exists() and CORPUS_PATH.exists():
            index = faiss.read_index(str(INDEX_PATH))
            with open(CORPUS_PATH, "r") as f:
                corpus = json.load(f)
            return index, corpus
    except Exception as e:
        st.error(f"Error loading index: {e}")
    return None, []

def save_index_and_corpus(index, corpus: List[Dict]):
    try:
        if index is not None:
            faiss.write_index(index, str(INDEX_PATH))
        with open(CORPUS_PATH, "w") as f:
            json.dump(corpus, f, indent=2)
    except Exception as e:
        st.error(f"Error saving index: {e}")

# --- KB processing -----------------------------------------------------------

def process_kb_files(silent: bool = False) -> List[Dict]:
    corpus: List[Dict] = []
    if not KB_DIR.exists():
        if not silent:
            st.error(f"KB_DIR does not exist: {KB_DIR}")
        return corpus

    files = list(KB_DIR.iterdir())
    if not silent and hasattr(st.session_state, "kb_loading") and st.session_state.kb_loading:
        st.info(f"Found {len(files)} files in KB directory. Processing... (this may take a while for {len(files)} files)")

    doc_files_processed = 0
    doc_files_failed = 0
    start_time = time.time()

    for file_idx, file_path in enumerate(files):
        if not file_path.is_file():
            continue
        try:
            suffix = file_path.suffix.lower()
            data = file_path.read_bytes()

            if suffix == ".docx":
                text = extract_text_from_docx_bytes(data)
            elif suffix == ".doc":
                text = extract_text_from_doc_bytes(data)
                if not text.strip():
                    doc_files_failed += 1
                    # Only log failures occasionally to avoid spam
                    if doc_files_failed <= 5 or doc_files_failed % 10 == 0:
                        print(f"⚠ No text extracted from .doc file: {file_path.name}")
                else:
                    doc_files_processed += 1
            elif suffix == ".pdf":
                text = extract_text_from_pdf_bytes(data)
            elif suffix in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
                if file_path.stat().st_size == 0:
                    continue
                text = extract_text_from_image_bytes(data)
            else:
                continue

            if text.strip():
                chunks = chunk_text(text)
                for i, chunk in enumerate(chunks):
                    corpus.append(
                        {
                            "text": chunk,
                            "source": file_path.name,
                            "chunk_id": i,
                            "file_type": suffix,
                        }
                    )

            if suffix in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"] and text.strip():
                try:
                    structured_items = parse_report_data_from_ocr(text, file_path.name)
                    for item in structured_items:
                        fields = [
                            f"Report: {item.get('report_type', 'Unknown')}",
                            f"Customer: {item.get('customer_name', '')}",
                            f"Contract: {item.get('contract', '')}",
                            f"Stock: {item.get('stock_number', '')}",
                            f"Make: {item.get('make', '')}",
                            f"Model: {item.get('model', '')}",
                            f"Type: {item.get('equipment_type', '')}",
                            f"Year: {item.get('year', '')}",
                            f"Serial: {item.get('serial', '')}",
                            f"Days Overdue: {item.get('days_overdue', '')}",
                            f"Date Out: {item.get('date_out', item.get('date_time_out', ''))}",
                            f"Expected Due: {item.get('expected_due', '')}",
                            f"Phone: {item.get('phone', '')}",
                            f"Location: {item.get('location', '')}",
                            f"Meter: {item.get('meter', '')}",
                        ]
                        searchable_text = " ".join(fields).strip()
                        corpus.append(
                            {
                                "text": searchable_text,
                                "source": file_path.name,
                                "chunk_id": len(corpus),
                                "file_type": suffix,
                                "content_type": "structured_data",
                                "structured_data": item,
                            }
                        )
                except Exception:
                    pass
            
            # Progress update every 50 files
            if not silent and (file_idx + 1) % 50 == 0:
                elapsed = time.time() - start_time
                rate = (file_idx + 1) / elapsed if elapsed > 0 else 0
                remaining = (len(files) - file_idx - 1) / rate if rate > 0 else 0
                print(f"Progress: {file_idx + 1}/{len(files)} files processed ({elapsed:.1f}s elapsed, ~{remaining:.0f}s remaining)")
                
        except Exception as e:
            if not silent:
                st.error(f"Error processing {file_path.name}: {e}")

    # Summary of .doc file processing
    elapsed_total = time.time() - start_time
    print(f"\n📊 Processing Summary:")
    print(f"   Total files: {len(files)}")
    print(f"   Total chunks created: {len(corpus)}")
    print(f"   .doc files processed: {doc_files_processed}")
    print(f"   .doc files failed: {doc_files_failed}")
    print(f"   Total time: {elapsed_total:.1f} seconds ({elapsed_total/60:.1f} minutes)")
    if doc_files_failed > 0:
        print(f"   💡 Tip: Install antiword or catdoc for better .doc support")

    validated: List[Dict] = []
    for item in corpus:
        tokens = estimate_tokens(item["text"])
        if tokens > 2000:
            for i, sub in enumerate(split_oversized_chunk(item["text"], 2000)):
                validated.append({**item, "text": sub, "chunk_id": f"{item['chunk_id']}_split_{i}"})
        else:
            validated.append(item)
    return validated

def get_conversation_context(messages: List[Dict], max_tokens: int = 2000) -> str:
    if not messages or len(messages) < 2:
        return ""
    recent = messages[-6:]
    context, current_tokens = [], 0
    for msg in reversed(recent):
        msg_text = f"{msg['role'].capitalize()}: {msg['content']}"
        tokens = estimate_tokens(msg_text)
        if current_tokens + tokens > max_tokens:
            break
        context.insert(0, msg_text)
        current_tokens += tokens
    return "\n".join(context)

# --- image handling ----------------------------------------------------------

def process_user_uploaded_image(image_bytes: bytes, query: str, model_name: str, project_id: str, location: str, credentials) -> str:
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = GenerativeModel(model_name)
        mime_type = "image/jpeg"
        if image_bytes.startswith(b"\x89PNG"):
            mime_type = "image/png"
        elif image_bytes.startswith(b"GIF"):
            mime_type = "image/gif"
        elif image_bytes.startswith(b"RIFF") and b"WEBP" in image_bytes[:12]:
            mime_type = "image/webp"
        image_part = Part.from_data(image_bytes, mime_type=mime_type)
        prompt = f"""You are an HBS assistant for NetView.

Analyze this image and answer the user's question: {query}

If unrelated to HBS/NetView, explain you specialize in that domain."""
        response = model.generate_content([prompt, image_part])
        return response.text if response.text else "I couldn't analyze the image. Please try again."
    except Exception as e:
        return f"Error analyzing image: {str(e)}"

# --- intent/sentiment --------------------------------------------------------

def analyze_user_sentiment_and_intent(query: str, conversation_context: str, model_name: str, project_id: str, location: str, credentials) -> Dict:
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = GenerativeModel(model_name)
        prompt = f"""Analyze the user's query and return JSON with intent, sentiment, context_relevance, escalation_needed, confidence, reasoning.

CONVERSATION CONTEXT:
{truncate_to_token_limit(conversation_context, 500)}

USER QUERY: {query}
"""
        response = model.generate_content(
            prompt,
            generation_config=GenerationConfig(temperature=0.1, max_output_tokens=300, top_p=0.8, top_k=40),
        )
        if response.text:
            try:
                return json.loads(response.text.strip())
            except json.JSONDecodeError:
                pass
    except Exception:
        pass
    return {
        "intent": "question",
        "sentiment": "neutral",
        "context_relevance": "new_topic",
        "escalation_needed": False,
        "confidence": 0.3,
        "reasoning": "Fallback classification",
    }

# --- response generation -----------------------------------------------------

def generate_semantic_response(query: str, context_chunks: List[Dict], user_analysis: Dict,
                               conversation_context: str, model_name: str, project_id: str,
                               location: str, credentials, deep_mode: bool) -> str:
    max_tokens = get_max_context_tokens(model_name)
    if deep_mode:
        context_chunks = cluster_chunks(context_chunks)
    context_text = build_optimized_context(context_chunks, max_tokens)
    context_section = ""
    if conversation_context:
        context_section = f"\nRECENT CONVERSATION CONTEXT:\n{truncate_to_token_limit(conversation_context, 2000)}\n"
    analysis_section = f"""
USER ANALYSIS:
- Intent: {user_analysis.get('intent', 'unknown')}
- Sentiment: {user_analysis.get('sentiment', 'neutral')}
- Context Relevance: {user_analysis.get('context_relevance', 'new_topic')}
- Escalation Needed: {user_analysis.get('escalation_needed', False)}
- Confidence: {user_analysis.get('confidence', 0):.2f}
- Reasoning: {user_analysis.get('reasoning', 'N/A')}
"""
    system_prompt = f"""You are an expert HBS NetView assistant. Provide accurate, actionable, and comprehensive answers.

SYSTEM CONTEXT:
You operate inside HBS Systems' NetView — a DMS for equipment dealerships.

{context_section}{analysis_section}

KNOWLEDGE BASE CONTEXT:
{context_text}

USER QUESTION: {query}

RESPONSE GUIDELINES:
1. **Direct Answer** first - provide a clear, direct response based on the context provided
2. **Detailed Explanation** - elaborate on the answer with relevant details from the context
3. **Steps/Procedures** - when present, list all steps in detail with specific field names, screen names, and values
4. **Examples** - include specific examples, field names, values, and dealership terminology when available
5. **Related Information** - add related tips, warnings, or important notes that might be helpful
6. **Multi-Document Integration** - it is possible that the complete answer to a question lies across several different documents. Each step in one document may relate to another in a different document. If that is the case, make sure you stitch together different parts of the answer from different documents and provide the complete, comprehensive answer to the user.
7. **Accuracy Priority** - accuracy takes priority, but provide as much detail as is available in the context

RULES:
- Answer based ONLY on the knowledge base context provided above
- If you can answer the question with the provided context, do so directly without mentioning what might be missing
- Only mention missing information if the user specifically asks about something and you need to clarify that specific aspect is not covered
- Do NOT include meta-commentary like "the documentation does not contain", "it's possible this exists but is not covered", or "the provided text may not have"
- Focus on providing helpful information from the context rather than explaining what's not there
- If information is genuinely critical and missing, simply state the limitation briefly (e.g., "To reverse a completed unit sale, please contact HBS Support as this process requires system access permissions.")
- Avoid hallucinations - only use information present in the context
- Be thorough and comprehensive - aim for 400-600 words for detailed answers
- Use bullets/numbered lists for clarity
- Include all relevant details, field names, screen references, and specific terminology
- When multiple documents contain related information, synthesize them into a cohesive answer"""
    
    if deep_mode:
        system_prompt += "\n\nMULTI-SOURCE INSTRUCTION: This is a complex query requiring information from multiple documents. Carefully combine information across all relevant documents, explain connections between different sources, and credit each source. Be comprehensive and detailed.\n"
    
    try:
        vertexai_init(project=project_id, location=location, credentials=credentials)
        model = GenerativeModel(model_name)
        response = model.generate_content(
            system_prompt,
            generation_config=GenerationConfig(temperature=0.1, max_output_tokens=8192, top_p=0.8, top_k=40),
        )
        answer = response.text if response.text else "I couldn't generate a response. Please try rephrasing your question."
        
        # Only add verification note if answer quality is poor, not for missing info
        verification = verify_answer_quality(query, answer, context_chunks, model_name, project_id, location, credentials)
        if verification.get("needs_improvement") and verification.get("support", 0) < 3:
            answer += "\n\n*Note: Some details may require confirmation with HBS Support.*"
        return answer
    except Exception as e:
        return f"Error generating response: {str(e)}"

# --- escalation --------------------------------------------------------------

def escalate_to_live_agent(query: str, conversation_context: str, user_analysis: Dict) -> str:
    summary = f"""
CONVERSATION SUMMARY FOR LIVE AGENT
===================================
USER QUESTION: {query}

CONTEXT:
{conversation_context}

USER ANALYSIS:
- Intent: {user_analysis.get('intent', 'unknown')}
- Sentiment: {user_analysis.get('sentiment', 'neutral')}
- Relevance: {user_analysis.get('context_relevance', 'new_topic')}
- Confidence: {user_analysis.get('confidence', 0):.2f}
- Reasoning: {user_analysis.get('reasoning', 'N/A')}
===================================
"""
    st.session_state.setdefault("escalation_requests", []).append(
        {
            "timestamp": len(st.session_state.messages),
            "query": query,
            "conversation_summary": summary,
            "user_analysis": user_analysis,
        }
    )
    esc_id = f"ESC-{len(st.session_state.messages):04d}"
    return "\n".join(
        [
            "I understand you need additional assistance. Let me connect you with an HBS Support Technician.",
            "",
            "**Connecting you with an HBS Support Technician now...**",
            "",
            f"Your question: {query}",
            "",
            "**What to expect:**",
            "- An HBS Support Technician will join the chat shortly",
            "- They can provide specialized assistance",
            "",
            f"**Reference ID:** {esc_id}",
            "",
            "Please hold while I connect you with a support technician...",
        ]
    )

# --- source summary helper ---------------------------------------------------

def summarize_sources(chunks: List[Dict]) -> List[Dict]:
    summaries: Dict[str, Dict[str, float]] = {}
    for chunk in chunks:
        source = chunk.get("source", "Unknown")
        entry = summaries.setdefault(source, {"source": source, "similarity": 0.0, "rerank": 0.0, "count": 0})
        entry["similarity"] = max(entry["similarity"], chunk.get("similarity_score", 0.0))
        entry["rerank"] = max(entry["rerank"], chunk.get("rerank_score", 0.0))
        entry["count"] += 1
    return sorted(summaries.values(), key=lambda x: (x["rerank"], x["similarity"]), reverse=True)

# --- main --------------------------------------------------------------------

def main():
    st.set_page_config(page_title="HBS Help Chatbot", page_icon="🤖", layout="wide")

    defaults = {
        "messages": [],
        "index": None,
        "corpus": [],
        "creds": None,
        "project_id": None,
        "location": None,
        "model_name": CANDIDATE_MODELS[0],
        "kb_loaded": False,
        "kb_loading": False,
        "escalation_requests": [],
        "last_kb_check": 0,
    }
    for key, value in defaults.items():
        st.session_state.setdefault(key, value)

    try:
        sa_info = json.loads(st.secrets["google"]["credentials_json"])
        st.session_state.creds = service_account.Credentials.from_service_account_info(sa_info)
        st.session_state.project_id = st.secrets["google"]["project"]
        st.session_state.location = st.secrets["google"].get("location", DEFAULT_LOCATION)
    except Exception as e:
        st.error(f"Error loading credentials: {e}")
        st.stop()

    @st.cache_resource
    def initialize_app():
        index, corpus = load_index_and_corpus()
        if index is not None and corpus:
            return index, corpus, True
        corpus = process_kb_files(silent=True)
        if not corpus:
            return None, [], False
        index, corpus = build_faiss_index(corpus, st.session_state.project_id, st.session_state.location, st.session_state.creds, silent=True)
        if index is not None:
            save_index_and_corpus(index, corpus)
            return index, corpus, True
        return None, [], False

    if not st.session_state.kb_loaded:
        st.session_state.kb_loading = True
        with st.spinner("Loading knowledge base... (Processing 600+ files - this will take 30-60 minutes)"):
            index, corpus, loaded = initialize_app()
            st.session_state.index = index
            st.session_state.corpus = corpus
            st.session_state.kb_loaded = loaded
            st.session_state.kb_loading = False

    if not st.session_state.get("polling_started", False):
        start_database_polling()
        st.session_state.polling_started = True

    with st.sidebar:
        st.header("HBS Help Chatbot")
        st.subheader("Model Settings")
        current_index = CANDIDATE_MODELS.index(st.session_state.model_name) if st.session_state.model_name in CANDIDATE_MODELS else 0
        st.session_state.model_name = st.selectbox(
            "Select Model",
            CANDIDATE_MODELS,
            index=current_index,
            key="model_select",
        )
        if st.session_state.escalation_requests:
            st.subheader("📞 Live Agent Requests")
            for i, req in enumerate(st.session_state.escalation_requests):
                with st.expander(f"Request #{i + 1} - {req['query'][:50]}..."):
                    st.write(f"**Query:** {req['query']}")
                    st.write(f"**Intent:** {req['user_analysis'].get('intent', 'unknown')}")
                    st.write(f"**Sentiment:** {req['user_analysis'].get('sentiment', 'unknown')}")
                    st.write(f"**Reference ID:** ESC-{req['timestamp']:04d}")
        if st.button("🔄 Rebuild Index"):
            INDEX_PATH.unlink(missing_ok=True)
            CORPUS_PATH.unlink(missing_ok=True)
            st.session_state.kb_loaded = False
            st.session_state.kb_loading = True
            st.cache_resource.clear()
            st.success("Cache cleared! Rebuilding knowledge base...")
            st.rerun()
        if st.button("🗑️ Clear Conversation"):
            st.session_state.messages = []
            st.rerun()

    st.title("HBS Help Chatbot")

    if not st.session_state.messages:
        st.info("Hi! How can I help you today?")

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])
            if "sources" in message and message["sources"]:
                display_sources = summarize_sources(message["sources"])
                with st.expander("📄 Sources"):
                    for src in display_sources:
                        st.write(f"📄 {src['source']} (chunks: {src['count']}, sim: {src['similarity']:.3f}, rerank: {src['rerank']:.3f})")

    upload_key = f"image_uploader_{len(st.session_state.messages)}"
    uploaded_image = st.file_uploader(
        "📷 Upload Image for Analysis",
        type=["png", "jpg", "jpeg", "webp", "bmp", "tiff"],
        key=upload_key,
    )

    if prompt := st.chat_input("Ask me anything about HBS systems..."):
        st.session_state.messages.append({"role": "user", "content": prompt})

        if uploaded_image is not None:
            with st.spinner("Analyzing your image..."):
                image_bytes = uploaded_image.read()
                response = process_user_uploaded_image(
                    image_bytes,
                    prompt,
                    st.session_state.model_name,
                    st.session_state.project_id,
                    st.session_state.location,
                    st.session_state.creds,
                )
                st.session_state.messages.append({"role": "assistant", "content": response, "timestamp": len(st.session_state.messages)})
            st.rerun()
        else:
            conversation_context = get_conversation_context(st.session_state.messages)
            with st.spinner("Understanding your request..."):
                user_analysis = analyze_user_sentiment_and_intent(
                    prompt,
                    conversation_context,
                    st.session_state.model_name,
                    st.session_state.project_id,
                    st.session_state.location,
                    st.session_state.creds,
                )

            if user_analysis.get("escalation_needed", False):
                response = escalate_to_live_agent(prompt, conversation_context, user_analysis)
                st.session_state.messages.append({"role": "assistant", "content": response, "timestamp": len(st.session_state.messages)})
            else:
                deep_mode = needs_deep_retrieval(prompt)
                with st.spinner("Searching knowledge base..."):
                    context_chunks = search_index(
                        prompt,
                        st.session_state.index,
                        st.session_state.corpus,
                        st.session_state.project_id,
                        st.session_state.location,
                        st.session_state.creds,
                        st.session_state.model_name,
                        deep_mode=deep_mode,
                        k=MAX_CHUNKS_FINAL,
                    )

                    response = generate_semantic_response(
                        prompt,
                        context_chunks,
                        user_analysis,
                        conversation_context,
                        st.session_state.model_name,
                        st.session_state.project_id,
                        st.session_state.location,
                        st.session_state.creds,
                        deep_mode=deep_mode,
                    )

                    st.session_state.messages.append(
                        {
                            "role": "assistant",
                            "content": response,
                            "sources": context_chunks,
                            "timestamp": len(st.session_state.messages),
                        }
                    )
        st.rerun()

if __name__ == "__main__":
    main()
